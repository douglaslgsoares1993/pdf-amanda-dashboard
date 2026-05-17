"""Geração dos relatórios TXT, DOCX, XLSX e exportações de busca."""

from __future__ import annotations

import hashlib
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from openpyxl import Workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


# Cores centralizadas para personalização futura.
COR_AZUL_ESCURO = "1F3864"
COR_AZUL_MEDIO = "2E74B5"
COR_AZUL_CLARO = "D6E4F0"
COR_AZUL_MUITO_CLARO = "F0F7FF"
COR_LINHA_ALTERNADA = "F7FBFF"
COR_CINZA_CLARO = "F2F2F2"
COR_CINZA_BORDA = "BFBFBF"
COR_CINZA_TEXTO = "666666"
COR_CINZA_PAGINA = "EEEEEE"
COR_TOTAL = "FFF2CC"
COR_BRANCO = "FFFFFF"


def cor_hex_para_rgb(hex_color: str):
    """Converte #RRGGBB/RRGGBB para RGBColor do python-docx."""
    from docx.shared import RGBColor

    cor = hex_color.strip().lstrip("#")
    return RGBColor(int(cor[0:2], 16), int(cor[2:4], 16), int(cor[4:6], 16))


def nome_base(arquivo: str) -> str:
    return Path(arquivo).stem.replace(" ", "_")


def nome_base_doc(doc: dict[str, Any]) -> str:
    """Nome de saída estável; usa hash curto quando há PDFs com mesmo nome."""
    if doc.get("saida_base"):
        return str(doc["saida_base"])
    caminho = str(doc.get("caminho", doc.get("arquivo", "")))
    base = nome_base(doc.get("arquivo", Path(caminho).name))
    if doc.get("precisa_hash"):
        curto = hashlib.sha1(caminho.encode("utf-8", errors="ignore")).hexdigest()[:4]
        return f"{base}_{curto}"
    return base


def garantir_saida(base_dir: str | Path) -> Path:
    saida = Path(base_dir) / "SAIDA"
    saida.mkdir(exist_ok=True)
    (saida / "erros.log").touch(exist_ok=True)
    return saida


def limpar_espacos(texto: str) -> str:
    """Evita blocos com muitas linhas em branco consecutivas."""
    return re.sub(r"\n{3,}", "\n\n", texto).strip() + "\n"


def _primeiro(campos: dict[str, list[str]], nome: str) -> str:
    vals = campos.get(nome, [])
    return vals[0] if vals else ""


def _valor(campos: dict[str, list[str]], nome: str, padrao: str = "—") -> str:
    valor = _primeiro(campos, nome)
    return valor if valor else padrao


def _juntar(campos: dict[str, list[str]], nome: str, padrao: str = "—") -> str:
    vals = [v for v in campos.get(nome, []) if v]
    return "; ".join(vals) if vals else padrao


def _categoria_do_campo(campo: str) -> str:
    mapa = {
        "Paciente / Nome": "Dados do Paciente",
        "Data de Nascimento": "Dados do Paciente",
        "CPF": "Dados do Paciente",
        "Data do Exame / Procedimento": "Dados do Procedimento",
        "Procedimento / Nome do Exame": "Dados do Procedimento",
        "CID": "Dados do Procedimento",
        "Código TUSS / CBHPM": "Dados do Procedimento",
        "Lateralidade": "Dados do Procedimento",
        "Médico Responsável": "Responsáveis",
        "CRM do médico": "Responsáveis",
        "Médico solicitante": "Responsáveis",
        "Hospital / Clínica / Instituição": "Responsáveis",
        "Convênio / Operadora / Plano": "Responsáveis",
        "Valores monetários (R$)": "Dados Financeiros",
        "Diagnóstico / Conclusão / Laudo": "Diagnóstico",
        "Número de Prontuário / Processo": "Dados do Paciente",
    }
    return mapa.get(campo, "Outros")


def _categorias_campos(campos: dict[str, list[str]]) -> list[tuple[str, list[tuple[str, str, str]]]]:
    return [
        ("DADOS DO PACIENTE", [
            ("Paciente", "Paciente / Nome", _valor(campos, "Paciente / Nome")),
            ("Nascimento", "Data de Nascimento", _valor(campos, "Data de Nascimento")),
            ("CPF", "CPF", _valor(campos, "CPF")),
            ("Prontuário", "Número de Prontuário / Processo", _valor(campos, "Número de Prontuário / Processo")),
        ]),
        ("DADOS DO PROCEDIMENTO", [
            ("Procedimento", "Procedimento / Nome do Exame", _valor(campos, "Procedimento / Nome do Exame")),
            ("CID", "CID", _juntar(campos, "CID")),
            ("Data do Exame", "Data do Exame / Procedimento", _valor(campos, "Data do Exame / Procedimento")),
            ("Lateralidade", "Lateralidade", _juntar(campos, "Lateralidade")),
            ("Código TUSS", "Código TUSS / CBHPM", _juntar(campos, "Código TUSS / CBHPM")),
        ]),
        ("RESPONSÁVEIS", [
            ("Médico", "Médico Responsável", _valor(campos, "Médico Responsável")),
            ("CRM", "CRM do médico", _valor(campos, "CRM do médico")),
            ("Solicitante", "Médico solicitante", _valor(campos, "Médico solicitante")),
            ("Hospital", "Hospital / Clínica / Instituição", _valor(campos, "Hospital / Clínica / Instituição")),
            ("Convênio", "Convênio / Operadora / Plano", _valor(campos, "Convênio / Operadora / Plano")),
        ]),
        ("DADOS FINANCEIROS", [
            ("Valor", "Valores monetários (R$)", _juntar(campos, "Valores monetários (R$)")),
        ]),
        ("DIAGNÓSTICO / CONCLUSÃO", [
            ("Diagnóstico", "Diagnóstico / Conclusão / Laudo", _juntar(campos, "Diagnóstico / Conclusão / Laudo")),
        ]),
    ]


def _ficha_resumo(doc: dict[str, Any]) -> dict[str, str]:
    campos = doc.get("campos", {})
    return {
        "DOCUMENTO": doc.get("arquivo", "—"),
        "PACIENTE": _valor(campos, "Paciente / Nome"),
        "EXAME": _valor(campos, "Procedimento / Nome do Exame"),
        "DATA": _valor(campos, "Data do Exame / Procedimento"),
        "MÉDICO": _valor(campos, "Médico Responsável"),
        "CRM": _valor(campos, "CRM do médico"),
        "CID": _juntar(campos, "CID"),
        "CONVÊNIO": _valor(campos, "Convênio / Operadora / Plano"),
        "VALOR": _juntar(campos, "Valores monetários (R$)"),
        "MOTOR": f"{doc.get('motor', '—')}   PÁGINAS: {doc.get('numero_paginas', '—')}",
    }


def _txt_ficha_resumo(doc: dict[str, Any]) -> str:
    ficha = _ficha_resumo(doc)
    largura = 68
    linhas = ["┌" + "─" * largura + "┐"]
    for campo, valor in ficha.items():
        texto = f"{campo:<9}: {valor}"
        linhas.append("│ " + texto[: largura - 2].ljust(largura - 2) + " │")
    linhas.append("└" + "─" * largura + "┘")
    return "\n".join(linhas)


def _txt_linha_campo(rotulo: str, valor: str, largura: int = 20) -> str:
    return f"{rotulo:<{largura}} {'.' * 12} {valor}"


def _txt_campos_categorizados(campos: dict[str, list[str]]) -> str:
    blocos: list[str] = ["CAMPOS EXTRAÍDOS"]
    for categoria, itens in _categorias_campos(campos):
        blocos.append(f"\n[{categoria}]")
        for rotulo, _, valor in itens:
            if categoria == "DIAGNÓSTICO / CONCLUSÃO":
                blocos.append(valor)
            else:
                blocos.append(_txt_linha_campo(rotulo, valor))
    return "\n".join(blocos)


def _txt_tabela(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    max_cols = max(len(row) for row in rows)
    normalizadas = [row + [""] * (max_cols - len(row)) for row in rows]
    larguras = [
        min(42, max(len(str(row[col] or "")) for row in normalizadas) + 2)
        for col in range(max_cols)
    ]
    sep = "─" * max(20, sum(larguras))
    linhas: list[str] = []
    for i, row in enumerate(normalizadas):
        valores = [str(row[col] or "").replace("\n", " ")[: larguras[col] - 1].ljust(larguras[col]) for col in range(max_cols)]
        linhas.append("".join(valores).rstrip())
        if i == 0:
            linhas.append(sep)
    linhas.append(sep)
    return "\n".join(linhas)


def _txt_analise(analise: dict[str, Any]) -> str:
    tipo = analise.get("tipo_documento", "—")
    resumo = analise.get("resumo", "—")
    procedimentos = analise.get("procedimentos", [])
    dados = analise.get("dados_clinicos", [])
    valores = analise.get("valores", [])
    pendencias = analise.get("pendencias", [])
    conclusao = analise.get("conclusao", "—")
    linhas = [
        "╔" + "═" * 62 + "╗",
        "║" + "ANÁLISE INTELIGENTE".center(62) + "║",
        "╠" + "═" * 62 + "╣",
        "║ " + f"Tipo: {tipo}"[:60].ljust(60) + " ║",
        "╠" + "═" * 62 + "╣",
        "RESUMO",
        str(resumo),
        "",
        "PROCEDIMENTOS",
    ]
    linhas.extend([f"• {item}" for item in procedimentos] or ["• —"])
    if dados:
        linhas.extend(["", "DADOS CLÍNICOS", *[f"• {item}" for item in dados]])
    if valores:
        linhas.extend(["", "VALORES", *[f"• {item}" for item in valores]])
    if pendencias:
        linhas.extend(["", "PENDÊNCIAS", *[f"• {item}" for item in pendencias]])
    linhas.extend(["", "CONCLUSÃO", str(conclusao), "╚" + "═" * 62 + "╝"])
    return "\n".join(linhas)


def gerar_texto_extraido(doc: dict[str, Any], saida: Path, reprocessar: bool = False) -> Path:
    caminho = saida / f"{nome_base_doc(doc)}_TEXTO.txt"
    if caminho.exists() and not reprocessar:
        return caminho
    linhas = []
    for pagina in doc.get("paginas", []):
        linhas.append(f"══════════════════════════  PÁGINA {pagina.get('page')}  ══════════════════════════")
        linhas.append(pagina.get("text", ""))
    caminho.write_text(limpar_espacos("\n\n".join(linhas)), encoding="utf-8-sig")
    return caminho


def gerar_relatorio_txt(doc: dict[str, Any], analise: dict[str, Any], saida: Path, reprocessar: bool = False) -> Path:
    caminho = saida / f"{nome_base_doc(doc)}_RELATORIO.txt"
    if caminho.exists() and not reprocessar:
        return caminho
    gerado = datetime.now().strftime("%d/%m/%Y %H:%M")
    linhas = [
        _txt_ficha_resumo(doc),
        "",
        f"Pasta de origem: {doc.get('origem', Path(doc.get('caminho', '')).parent)}",
        f"Tamanho do arquivo: {doc.get('tamanho_arquivo')} bytes",
        f"Score de extração: {doc.get('score')}",
        "",
        _txt_campos_categorizados(doc.get("campos", {})),
        "",
        "METADADOS DO ARQUIVO",
    ]
    meta = doc.get("metadados", {})
    for chave, valor in meta.items():
        if valor:
            linhas.append(f"{chave:<18} {valor}")
    if not any(meta.values()):
        linhas.append("—")

    linhas.extend(["", _txt_analise(analise), "", "TABELAS EXTRAÍDAS"])
    tabelas = doc.get("tabelas", [])
    if not tabelas:
        linhas.append("Nenhuma tabela identificada.")
    for tabela in tabelas:
        linhas.append(f"\nTabela {tabela.get('table_index')} — Página {tabela.get('page')}")
        linhas.append(_txt_tabela(tabela.get("rows", [])))

    linhas.extend(["", "TEXTO EXTRAÍDO DO DOCUMENTO"])
    for pagina in doc.get("paginas", []):
        linhas.append(f"\n══════════════════════════  PÁGINA {pagina.get('page')}  ══════════════════════════")
        linhas.append(pagina.get("text", ""))
    linhas.append(f"\n── Gerado em {gerado} | Motor: {doc.get('motor', '—')} | PDF Amanda System ──")
    caminho.write_text(limpar_espacos("\n".join(linhas)), encoding="utf-8-sig")
    return caminho


# Helpers DOCX.
def _docx_set_cell_shading(cell, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def _docx_set_cell_borders(cell, color: str = COR_CINZA_BORDA, size: str = "6", left_size: str | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), left_size if edge == "left" and left_size else size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _docx_set_cell_width(cell, width_twips: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def _docx_format_paragraph(paragraph, size: int = 11, bold: bool = False, color: str | None = None, font: str = "Arial") -> None:
    for run in paragraph.runs:
        run.font.name = font
        run.font.size = _pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = cor_hex_para_rgb(color)


def _pt(value: int | float):
    from docx.shared import Pt

    return Pt(value)


def _docx_add_page_number(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph.add_run("Página ")
    for field in ("PAGE", "NUMPAGES"):
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        paragraph._p.append(fld_begin)
        paragraph._p.append(instr)
        paragraph._p.append(fld_end)
        if field == "PAGE":
            paragraph.add_run(" de ")


def _docx_paragraph_bottom_border(paragraph, color: str = COR_CINZA_BORDA) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _docx_configurar_estilos(document) -> None:
    from docx.enum.style import WD_STYLE_TYPE

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = _pt(11)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 13, COR_AZUL_ESCURO),
        ("Heading 2", 11, COR_AZUL_MEDIO),
        ("Heading 3", 10, COR_AZUL_MEDIO),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = _pt(size)
        style.font.bold = True
        style.font.color.rgb = cor_hex_para_rgb(color)
        style.paragraph_format.space_before = _pt(12 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = _pt(6 if name == "Heading 1" else 3)

    if "Marcador de Página PDF" not in styles:
        style = styles.add_style("Marcador de Página PDF", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style.font.size = _pt(9)
        style.font.color.rgb = cor_hex_para_rgb(COR_CINZA_TEXTO)
        style.paragraph_format.space_before = _pt(8)
        style.paragraph_format.space_after = _pt(4)


def _docx_configurar_cabecalho_rodape(document, doc: dict[str, Any], data: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    section = document.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    p.text = f"{doc.get('arquivo', '')}\t{data}"
    p.style = document.styles["Normal"]
    p.paragraph_format.tab_stops.add_tab_stop(_cm(16))
    _docx_format_paragraph(p, size=8, color=COR_CINZA_TEXTO)
    _docx_paragraph_bottom_border(p)

    footer = section.footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_add_page_number(p_footer)
    p_footer.add_run("\tPDF Amanda System")
    p_footer.paragraph_format.tab_stops.add_tab_stop(_cm(15))
    _docx_format_paragraph(p_footer, size=8, color=COR_CINZA_TEXTO)


def _cm(value: int | float):
    from docx.shared import Cm

    return Cm(value)


def _docx_add_capa(document, doc: dict[str, Any], data: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    titulo = document.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(doc.get("arquivo", "Relatório PDF"))
    run.font.name = "Arial"
    run.font.size = _pt(18)
    run.font.bold = True
    run.font.color.rgb = cor_hex_para_rgb(COR_AZUL_ESCURO)
    _docx_paragraph_bottom_border(titulo, COR_AZUL_MEDIO)

    tabela = document.add_table(rows=0, cols=2)
    tabela.autofit = False
    for campo, valor in _ficha_resumo(doc).items():
        row = tabela.add_row().cells
        row[0].text = campo.title()
        row[1].text = valor
        _docx_set_cell_width(row[0], 2200)
        _docx_set_cell_width(row[1], 6200)
        _docx_set_cell_borders(row[0], COR_CINZA_BORDA, "4")
        _docx_set_cell_borders(row[1], COR_CINZA_BORDA, "4")
        for paragraph in row[0].paragraphs:
            _docx_format_paragraph(paragraph, size=10, bold=True, color=COR_AZUL_ESCURO)
        for paragraph in row[1].paragraphs:
            _docx_format_paragraph(paragraph, size=10)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Processado em {data} | Motor: {doc.get('motor', '—')}")
    _docx_format_paragraph(p, size=9, color=COR_CINZA_TEXTO)
    document.add_page_break()


def _docx_add_metadados(document, meta: dict[str, str]) -> None:
    dados = [(k, v) for k, v in meta.items() if v]
    document.add_heading("Metadados do Arquivo", level=1)
    if not dados:
        document.add_paragraph("—")
        return
    tabela = document.add_table(rows=0, cols=2)
    for chave, valor in dados:
        row = tabela.add_row().cells
        row[0].text = str(chave)
        row[1].text = str(valor)
        for cell in row:
            _docx_set_cell_borders(cell, COR_BRANCO, "0")
            for paragraph in cell.paragraphs:
                _docx_format_paragraph(paragraph, size=9, color=COR_CINZA_TEXTO)


def _docx_add_campos(document, campos: dict[str, list[str]]) -> None:
    document.add_heading("Campos Extraídos", level=1)
    tabela = document.add_table(rows=0, cols=2)
    tabela.style = "Table Grid"
    alternar = False
    for categoria, itens in _categorias_campos(campos):
        cells = tabela.add_row().cells
        cells[0].merge(cells[1])
        cells[0].text = categoria.title()
        _docx_set_cell_shading(cells[0], COR_AZUL_CLARO)
        _docx_set_cell_borders(cells[0])
        _docx_format_paragraph(cells[0].paragraphs[0], size=11, bold=True, color=COR_AZUL_ESCURO)
        for rotulo, _, valor in itens:
            row = tabela.add_row().cells
            row[0].text = rotulo
            row[1].text = "" if valor == "—" else valor
            _docx_set_cell_width(row[0], 2835)
            fill = COR_LINHA_ALTERNADA if alternar else COR_BRANCO
            alternar = not alternar
            for cell in row:
                _docx_set_cell_shading(cell, fill)
                _docx_set_cell_borders(cell)
            _docx_format_paragraph(row[0].paragraphs[0], size=10, bold=True)
            if valor == "—":
                row[1].text = "—"
                for run in row[1].paragraphs[0].runs:
                    run.italic = True
                    run.font.color.rgb = cor_hex_para_rgb(COR_CINZA_TEXTO)
            _docx_format_paragraph(row[1].paragraphs[0], size=10)


def _docx_add_analise(document, analise: dict[str, Any]) -> None:
    document.add_heading("Análise Inteligente", level=1)
    box = document.add_table(rows=1, cols=1)
    cell = box.rows[0].cells[0]
    _docx_set_cell_shading(cell, COR_AZUL_MUITO_CLARO)
    _docx_set_cell_borders(cell, COR_AZUL_MEDIO, size="4", left_size="18")
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run("Análise Inteligente").bold = True
    _docx_format_paragraph(p, size=12, bold=True, color=COR_AZUL_ESCURO)
    for titulo, chave in [
        ("Tipo", "tipo_documento"),
        ("Resumo", "resumo"),
        ("Procedimentos", "procedimentos"),
        ("Dados Clínicos", "dados_clinicos"),
        ("Valores", "valores"),
        ("Pendências", "pendencias"),
        ("Conclusão", "conclusao"),
    ]:
        sub = cell.add_paragraph()
        sub.add_run(titulo).bold = True
        _docx_format_paragraph(sub, size=10, bold=True, color=COR_AZUL_MEDIO)
        valor = analise.get(chave, "")
        if isinstance(valor, list):
            if valor:
                for item in valor:
                    p_item = cell.add_paragraph(f"• {item}")
                    _docx_format_paragraph(p_item, size=10)
            else:
                p_item = cell.add_paragraph("—")
                _docx_format_paragraph(p_item, size=10, color=COR_CINZA_TEXTO)
        else:
            p_valor = cell.add_paragraph(str(valor or "—"))
            _docx_format_paragraph(p_valor, size=10)


def _docx_add_tabelas_pdf(document, tabelas: list[dict[str, Any]]) -> None:
    document.add_heading("Tabelas Extraídas", level=1)
    if not tabelas:
        document.add_paragraph("Nenhuma tabela identificada.")
        return
    for tabela in tabelas:
        rows = tabela.get("rows", [])
        document.add_heading(f"Tabela {tabela.get('table_index')} — Página {tabela.get('page')}", level=2)
        if not rows:
            continue
        max_cols = max(len(r) for r in rows)
        tdoc = document.add_table(rows=0, cols=max_cols)
        tdoc.style = "Table Grid"
        for ri, linha in enumerate(rows):
            cells = tdoc.add_row().cells
            is_total = any("total" in str(cel or "").lower() for cel in linha)
            for ci in range(max_cols):
                cells[ci].text = str(linha[ci] if ci < len(linha) else "")
                _docx_set_cell_borders(cells[ci])
                if ri == 0:
                    _docx_set_cell_shading(cells[ci], COR_AZUL_ESCURO)
                    _docx_format_paragraph(cells[ci].paragraphs[0], size=9, bold=True, color=COR_BRANCO)
                elif is_total:
                    _docx_set_cell_shading(cells[ci], COR_TOTAL)
                    _docx_format_paragraph(cells[ci].paragraphs[0], size=9, bold=True)
                else:
                    _docx_set_cell_shading(cells[ci], COR_CINZA_CLARO if ri % 2 == 0 else COR_BRANCO)
                    _docx_format_paragraph(cells[ci].paragraphs[0], size=9)


def _docx_add_texto_completo(document, paginas: list[dict[str, Any]]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document.add_heading("Texto Extraído do Documento", level=1)
    for pagina in paginas:
        p = document.add_paragraph(style="Marcador de Página PDF")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"PÁGINA {pagina.get('page')}")
        run.font.name = "Arial"
        run.font.size = _pt(9)
        run.font.color.rgb = cor_hex_para_rgb(COR_CINZA_TEXTO)
        p_format = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), COR_CINZA_PAGINA)
        p_format.append(shd)
        texto = document.add_paragraph()
        texto.add_run(pagina.get("text", ""))
        _docx_format_paragraph(texto, size=9, font="Courier New")


def gerar_relatorio_docx(doc: dict[str, Any], analise: dict[str, Any], saida: Path, reprocessar: bool = False) -> Path:
    from docx import Document

    caminho = saida / f"{nome_base_doc(doc)}_RELATORIO.docx"
    if caminho.exists() and not reprocessar:
        return caminho

    data = datetime.now().strftime("%d/%m/%Y %H:%M")
    document = Document()
    _docx_configurar_estilos(document)
    _docx_configurar_cabecalho_rodape(document, doc, data)
    _docx_add_capa(document, doc, data)
    _docx_add_campos(document, doc.get("campos", {}))
    _docx_add_analise(document, analise)
    _docx_add_metadados(document, doc.get("metadados", {}))
    _docx_add_tabelas_pdf(document, doc.get("tabelas", []))
    document.add_page_break()
    _docx_add_texto_completo(document, doc.get("paginas", []))
    document.save(caminho)
    return caminho


# Helpers XLSX.
def ajustar_largura_colunas(ws) -> None:
    for col in ws.columns:
        letra = get_column_letter(col[0].column)
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[letra].width = min(max(max_length + 4, 12), 60)


def _xlsx_border() -> Border:
    side = Side(style="thin", color=COR_CINZA_BORDA)
    return Border(left=side, right=side, top=side, bottom=side)


def _xlsx_header(ws, row: int = 1, max_col: int | None = None) -> None:
    max_col = max_col or ws.max_column
    for cell in ws[row][:max_col]:
        cell.font = Font(bold=True, color=COR_BRANCO)
        cell.fill = PatternFill("solid", fgColor=COR_AZUL_ESCURO)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = _xlsx_border()
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def _xlsx_estilo_dados(ws) -> None:
    for row in ws.iter_rows():
        for cell in row:
            cell.font = cell.font.copy(name="Calibri", size=11)
            if cell.value is not None:
                cell.border = _xlsx_border()
                cell.alignment = Alignment(vertical="top", wrap_text=False)
    ws.sheet_view.zoomScale = 90


def _parse_valor_monetario(valor: str) -> float | None:
    if not valor or valor == "—":
        return None
    match = re.search(r"R\$\s*([\d.]+,\d{2})", valor)
    if not match:
        return None
    return float(match.group(1).replace(".", "").replace(",", "."))


def _parse_data(valor: str):
    if not valor or valor == "—":
        return ""
    match = re.search(r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}", valor)
    if not match:
        return valor
    texto = match.group(0)
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%d/%m/%y", "%d-%m-%y"):
        try:
            return datetime.strptime(texto, fmt)
        except ValueError:
            pass
    return valor


def _add_separador_arquivo(ws, row_idx: int, arquivo: str, cols: int) -> int:
    ws.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=cols)
    cell = ws.cell(row_idx, 1, arquivo)
    cell.fill = PatternFill("solid", fgColor=COR_AZUL_CLARO)
    cell.font = Font(bold=True, color=COR_AZUL_ESCURO)
    cell.border = _xlsx_border()
    return row_idx + 1


def gerar_consolidado(resultados: list[dict[str, Any]], analises: dict[str, dict[str, Any]], saida: Path) -> Path:
    caminho = saida / "CONSOLIDADO.xlsx"
    wb = Workbook()
    wb.calculation.fullCalcOnLoad = True

    ws = wb.active
    ws.title = "Resumo Geral"
    headers = [
        "Arquivo", "Pasta de Origem", "Páginas", "Motor de Extração", "Data de Processamento",
        "Paciente", "Data do Exame", "Data de Nascimento", "CPF", "Médico Responsável", "CRM",
        "Médico Solicitante", "Hospital / Clínica", "Convênio / Operadora", "Procedimento Principal",
        "Código TUSS", "CID", "Lateralidade", "Valor Total (R$)",
        "Diagnóstico / Conclusão (texto completo)", "Resumo IA (se disponível)", "Tipo de Documento (IA)",
        "Pendências (IA)",
    ]
    ws.append(headers)
    agora = datetime.now()
    for idx, doc in enumerate(resultados, start=2):
        campos = doc.get("campos", {})
        analise = analises.get(doc["arquivo"], {})
        valor_total = _parse_valor_monetario(_juntar(campos, "Valores monetários (R$)", ""))
        ws.append([
            doc["arquivo"],
            doc.get("origem", str(Path(doc.get("caminho", "")).parent)),
            doc.get("numero_paginas", 0),
            doc.get("motor", ""),
            agora,
            _primeiro(campos, "Paciente / Nome"),
            _parse_data(_primeiro(campos, "Data do Exame / Procedimento")),
            _parse_data(_primeiro(campos, "Data de Nascimento")),
            _primeiro(campos, "CPF"),
            _primeiro(campos, "Médico Responsável"),
            _primeiro(campos, "CRM do médico"),
            _primeiro(campos, "Médico solicitante"),
            _primeiro(campos, "Hospital / Clínica / Instituição"),
            _primeiro(campos, "Convênio / Operadora / Plano"),
            _primeiro(campos, "Procedimento / Nome do Exame"),
            _juntar(campos, "Código TUSS / CBHPM", ""),
            _juntar(campos, "CID", ""),
            _juntar(campos, "Lateralidade", ""),
            valor_total,
            _juntar(campos, "Diagnóstico / Conclusão / Laudo", ""),
            analise.get("resumo", ""),
            analise.get("tipo_documento", ""),
            "; ".join(analise.get("pendencias", []) or []),
        ])
        fill = PatternFill("solid", fgColor=COR_LINHA_ALTERNADA if idx % 2 == 0 else COR_BRANCO)
        for cell in ws[idx]:
            cell.fill = fill
    _xlsx_header(ws)
    for row in ws.iter_rows(min_row=2):
        row[4].number_format = "DD/MM/YYYY"
        row[6].number_format = "DD/MM/YYYY"
        row[7].number_format = "DD/MM/YYYY"
        row[18].number_format = 'R$ #,##0.00'
        for idx in (20, 21, 23):
            row[idx - 1].alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[row[0].row].height = 60
    ws.column_dimensions["T"].width = 60
    ws.column_dimensions["U"].width = 60
    ws.column_dimensions["W"].width = 60

    ws_campos = wb.create_sheet("Campos Extraídos Detalhados")
    ws_campos.append(["Arquivo", "Categoria do Campo", "Campo", "Valor Encontrado", "Ocorrência", "Página onde foi encontrado"])
    row_idx = 2
    for doc in resultados:
        row_idx = _add_separador_arquivo(ws_campos, row_idx, doc["arquivo"], 6)
        for campo, valores in doc.get("campos", {}).items():
            categoria = _categoria_do_campo(campo)
            if valores:
                for i, valor in enumerate(valores, start=1):
                    ws_campos.append([doc["arquivo"], categoria, campo, valor, f"{i}ª", ""])
                    row_idx += 1
            else:
                ws_campos.append([doc["arquivo"], categoria, campo, "", "", ""])
                row_idx += 1
    _xlsx_header(ws_campos)

    ws_tab = wb.create_sheet("Tabelas dos Documentos")
    tem_tabela = any(doc.get("tabelas") for doc in resultados)
    if not tem_tabela:
        ws_tab["A1"] = "Nenhuma tabela foi identificada nos documentos processados."
        ws_tab["A1"].font = Font(bold=True, color=COR_AZUL_ESCURO)
    else:
        row_idx = 1
        for doc in resultados:
            for tabela in doc.get("tabelas", []):
                rows = tabela.get("rows", [])
                max_cols = max((len(r) for r in rows), default=1)
                ws_tab.merge_cells(start_row=row_idx, start_column=1, end_row=row_idx, end_column=max_cols)
                title = ws_tab.cell(row_idx, 1, f"Tabela {tabela.get('table_index')} — Página {tabela.get('page')} — {doc['arquivo']}")
                title.fill = PatternFill("solid", fgColor=COR_AZUL_MEDIO)
                title.font = Font(bold=True, color=COR_BRANCO)
                row_idx += 1
                for ri, linha in enumerate(rows):
                    is_total = any("total" in str(c or "").lower() for c in linha)
                    for ci, valor in enumerate(linha, start=1):
                        cell = ws_tab.cell(row_idx, ci, valor)
                        if ri == 0:
                            cell.fill = PatternFill("solid", fgColor=COR_AZUL_CLARO)
                            cell.font = Font(bold=True)
                        elif is_total:
                            cell.fill = PatternFill("solid", fgColor=COR_TOTAL)
                            cell.font = Font(bold=True)
                        else:
                            cell.fill = PatternFill("solid", fgColor=COR_LINHA_ALTERNADA if ri % 2 == 1 else COR_BRANCO)
                    row_idx += 1
                row_idx += 2

    ws_texto = wb.create_sheet("Texto Completo")
    ws_texto.append(["Arquivo", "Página", "Texto da Página"])
    row_idx = 2
    for doc in resultados:
        row_idx = _add_separador_arquivo(ws_texto, row_idx, doc["arquivo"], 3)
        for pagina in doc.get("paginas", []):
            ws_texto.append([doc["arquivo"], pagina.get("page"), pagina.get("text", "")])
            ws_texto.cell(row_idx, 3).alignment = Alignment(wrap_text=True, vertical="top")
            ws_texto.row_dimensions[row_idx].height = 80
            row_idx += 1
    _xlsx_header(ws_texto)
    ws_texto.column_dimensions["C"].width = 120

    ws_busca = wb.create_sheet("Busca")
    ws_busca.append(["Termo Buscado", "Arquivo", "Página", "Trecho Encontrado", "Score de Relevância", "Data da Busca"])
    ws_busca["A1"].comment = Comment(
        "Esta aba é preenchida automaticamente quando você exporta resultados de busca pelo menu do sistema.",
        "PDF Amanda System",
    )
    _xlsx_header(ws_busca)

    for sheet in wb.worksheets:
        _xlsx_estilo_dados(sheet)
        ajustar_largura_colunas(sheet)
        sheet.sheet_view.zoomScale = 90
        if sheet.max_row >= 1 and sheet.max_column >= 1:
            sheet.auto_filter.ref = sheet.dimensions
    ws.column_dimensions["T"].width = 60
    ws.column_dimensions["U"].width = 60
    ws.column_dimensions["W"].width = 60
    ws_texto.column_dimensions["C"].width = 120
    wb.active = wb.sheetnames.index("Resumo Geral")
    wb.save(caminho)
    return caminho


def exportar_busca(resultados: list[dict[str, Any]], termo: str, formato: str, saida: Path) -> Path:
    from buscador import formatar_resultados, limpar_nome_arquivo

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = saida / f"BUSCA_{limpar_nome_arquivo(termo)}_{timestamp}.{formato.lower()}"
    if formato.lower() == "txt":
        base.write_text(formatar_resultados(resultados, limite=100000), encoding="utf-8-sig")
    elif formato.lower() == "docx":
        from docx import Document

        doc = Document()
        _docx_configurar_estilos(doc)
        doc.add_heading(f"Resultado de Busca: {termo}", level=1)
        for item in resultados:
            doc.add_paragraph(
                f"{item.get('arquivo', '')} | {item.get('caminho', item.get('origem', ''))} | "
                f"página {item.get('pagina', '')} | score {item.get('score', '')}"
            )
            doc.add_paragraph(str(item.get("contexto", "")))
        doc.save(base)
    elif formato.lower() == "xlsx":
        wb = Workbook()
        ws = wb.active
        ws.title = "Busca"
        ws.append(["Tipo", "Arquivo", "Caminho", "Origem", "Página", "Linha", "Coluna", "Campo", "Score", "Contexto"])
        for item in resultados:
            ws.append([
                item.get("tipo", ""), item.get("arquivo", ""), item.get("caminho", ""), item.get("origem", ""), item.get("pagina", ""),
                item.get("linha", ""), item.get("coluna", ""), item.get("campo", ""),
                item.get("score", ""), item.get("contexto", ""),
            ])
        _xlsx_header(ws)
        _xlsx_estilo_dados(ws)
        ajustar_largura_colunas(ws)
        wb.save(base)
    else:
        raise ValueError("Formato inválido. Use txt, docx ou xlsx.")
    return base
