"""
PDF_CONVERTER.PY - Conversor e Extrator Inteligente de PDF
-----------------------------------------------------------
Converte PDF para TXT e DOCX com cascata de motores de leitura:
  1. pdfplumber  - texto nativo + tabelas
  2. pymupdf     - extração avançada com layout
  3. pypdf       - fallback leve
  4. OCR         - pytesseract para PDFs escaneados (último recurso)

Extrator inteligente: identifica campos comuns em laudos, relatórios
e documentos sem depender de IA.

Uso: python pdf_converter.py
     (o script pergunta o caminho interativamente)
"""

import os
import sys
import re
import unicodedata
from pathlib import Path
from datetime import datetime

# ──────────────────────────────────────────────────────
# Imports opcionais com verificação
# ──────────────────────────────────────────────────────
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import fitz  # pymupdf
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import pytesseract
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ──────────────────────────────────────────────────────
# Utilitários
# ──────────────────────────────────────────────────────

SEPARADOR = "─" * 70

def limpar_texto(texto: str) -> str:
    """Remove lixo de codificação, normaliza espaços e quebras."""
    if not texto:
        return ""
    # Normaliza unicode
    texto = unicodedata.normalize("NFC", texto)
    # Remove caracteres de controle exceto \n e \t
    texto = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', texto)
    # Normaliza múltiplas quebras de linha
    texto = re.sub(r'\n{4,}', '\n\n\n', texto)
    # Remove espaços no início/fim de cada linha
    linhas = [l.rstrip() for l in texto.splitlines()]
    texto = '\n'.join(linhas)
    return texto.strip()


def avaliar_qualidade(texto: str) -> float:
    """
    Retorna score 0.0-1.0 indicando qualidade do texto extraído.
    Critérios: proporção de caracteres alfanuméricos, ausência de lixo.
    """
    if not texto or len(texto) < 50:
        return 0.0
    total = len(texto)
    alfanum = sum(1 for c in texto if c.isalnum() or c in ' \n,.;:()-/')
    proporcao = alfanum / total
    # Penaliza se tiver muitos caracteres estranhos
    lixo = sum(1 for c in texto if ord(c) > 65000)
    penalidade = lixo / total if total > 0 else 0
    return max(0.0, min(1.0, proporcao - penalidade * 2))


def normalizar_caminho(caminho_raw: str) -> Path:
    """
    Aceita caminhos Windows (C:\\...), WSL (/mnt/c/...), Unix, com ou sem
    aspas, com espaços, com OneDrive, etc.
    """
    # Remove aspas externas se houver
    caminho = caminho_raw.strip().strip('"').strip("'")

    # Converte separadores Windows para o sistema atual se necessário
    # (em Windows puro, Path() já lida com ambos)
    if sys.platform != 'win32':
        # Converte C:\... para /mnt/c/... se estiver em WSL
        match = re.match(r'^([A-Za-z]):\\(.*)', caminho)
        if match:
            drive = match.group(1).lower()
            resto = match.group(2).replace('\\', '/')
            caminho = f'/mnt/{drive}/{resto}'
        else:
            caminho = caminho.replace('\\', '/')

    return Path(caminho)


# ──────────────────────────────────────────────────────
# Motores de extração
# ──────────────────────────────────────────────────────

def extrair_pdfplumber(pdf_path: Path) -> tuple[str, list]:
    """Motor 1: pdfplumber - melhor para texto nativo e tabelas."""
    texto_total = []
    tabelas_total = []
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                t = page.extract_text()
                if t:
                    texto_total.append(f"[Página {i}]\n{t}")
                # Extrai tabelas
                tabelas = page.extract_tables()
                for j, tabela in enumerate(tabelas, 1):
                    if tabela:
                        tabelas_total.append((i, j, tabela))
    except Exception as e:
        return "", []
    return '\n\n'.join(texto_total), tabelas_total


def extrair_pymupdf(pdf_path: Path) -> str:
    """Motor 2: pymupdf - ótimo para PDFs médicos/complexos com layout."""
    texto_total = []
    try:
        doc = fitz.open(str(pdf_path))
        for i, page in enumerate(doc, 1):
            # Extração com preservação de blocos
            blocos = page.get_text("blocks")
            blocos_ordenados = sorted(blocos, key=lambda b: (b[1], b[0]))
            partes = []
            for b in blocos_ordenados:
                if b[6] == 0:  # tipo 0 = texto
                    partes.append(b[4].strip())
            texto_pagina = '\n'.join(p for p in partes if p)
            if texto_pagina:
                texto_total.append(f"[Página {i}]\n{texto_pagina}")
        doc.close()
    except Exception as e:
        return ""
    return '\n\n'.join(texto_total)


def extrair_pymupdf_detalhado(pdf_path: Path) -> str:
    """Motor 2b: pymupdf modo dict - captura estrutura de fonte/tamanho."""
    texto_total = []
    try:
        doc = fitz.open(str(pdf_path))
        for i, page in enumerate(doc, 1):
            dicionario = page.get_text("dict")
            linhas_pagina = []
            for bloco in dicionario.get("blocks", []):
                if bloco.get("type") == 0:
                    for linha in bloco.get("lines", []):
                        spans_texto = []
                        for span in linha.get("spans", []):
                            t = span.get("text", "").strip()
                            if t:
                                spans_texto.append(t)
                        if spans_texto:
                            linhas_pagina.append(" ".join(spans_texto))
            if linhas_pagina:
                texto_total.append(f"[Página {i}]\n" + "\n".join(linhas_pagina))
        doc.close()
    except Exception:
        return ""
    return '\n\n'.join(texto_total)


def extrair_pypdf(pdf_path: Path) -> str:
    """Motor 3: pypdf - fallback leve."""
    texto_total = []
    try:
        reader = PdfReader(str(pdf_path))
        for i, page in enumerate(reader.pages, 1):
            t = page.extract_text()
            if t:
                texto_total.append(f"[Página {i}]\n{t}")
    except Exception:
        return ""
    return '\n\n'.join(texto_total)


def extrair_ocr(pdf_path: Path, dpi: int = 300) -> str:
    """Motor 4: OCR via pytesseract - para PDFs escaneados."""
    texto_total = []
    try:
        print("    [OCR] Convertendo páginas para imagem...")
        imagens = convert_from_path(str(pdf_path), dpi=dpi)
        for i, img in enumerate(imagens, 1):
            print(f"    [OCR] Processando página {i}/{len(imagens)}...")
            config = "--oem 3 --psm 6 -l por+eng"
            t = pytesseract.image_to_string(img, config=config)
            if t.strip():
                texto_total.append(f"[Página {i}]\n{t}")
    except Exception as e:
        print(f"    [OCR] Erro: {e}")
        return ""
    return '\n\n'.join(texto_total)


def extrair_metadados(pdf_path: Path) -> dict:
    """Extrai metadados do PDF (autor, título, data de criação, etc.)."""
    meta = {}
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(str(pdf_path))
            m = doc.metadata
            doc.close()
            for k, v in m.items():
                if v:
                    meta[k] = v
        except Exception:
            pass
    if not meta and HAS_PYPDF:
        try:
            reader = PdfReader(str(pdf_path))
            m = reader.metadata
            if m:
                for k in m:
                    v = m[k]
                    if v:
                        meta[k.lstrip('/')] = str(v)
        except Exception:
            pass
    return meta


# ──────────────────────────────────────────────────────
# Cascata principal
# ──────────────────────────────────────────────────────

def extrair_texto_cascata(pdf_path: Path) -> tuple[str, list, dict, str]:
    """
    Tenta extrair texto usando cascata de motores.
    Retorna: (texto, tabelas, metadados, motor_usado)
    """
    tabelas = []
    meta = extrair_metadados(pdf_path)
    LIMIAR = 0.45  # qualidade mínima aceitável

    print(f"\n  Tentando motor 1: pdfplumber...")
    if HAS_PDFPLUMBER:
        texto, tabelas = extrair_pdfplumber(pdf_path)
        texto = limpar_texto(texto)
        score = avaliar_qualidade(texto)
        print(f"    Score de qualidade: {score:.2f}")
        if score >= LIMIAR:
            print(f"    ✔ pdfplumber suficiente.")
            return texto, tabelas, meta, "pdfplumber"
    else:
        print("    pdfplumber não disponível.")

    print(f"  Tentando motor 2: pymupdf (blocos)...")
    if HAS_PYMUPDF:
        texto = limpar_texto(extrair_pymupdf(pdf_path))
        score = avaliar_qualidade(texto)
        print(f"    Score de qualidade: {score:.2f}")
        if score >= LIMIAR:
            print(f"    ✔ pymupdf (blocos) suficiente.")
            return texto, tabelas, meta, "pymupdf"

        print(f"  Tentando motor 2b: pymupdf (dict)...")
        texto = limpar_texto(extrair_pymupdf_detalhado(pdf_path))
        score = avaliar_qualidade(texto)
        print(f"    Score de qualidade: {score:.2f}")
        if score >= LIMIAR:
            print(f"    ✔ pymupdf (dict) suficiente.")
            return texto, tabelas, meta, "pymupdf-dict"
    else:
        print("    pymupdf não disponível.")

    print(f"  Tentando motor 3: pypdf...")
    if HAS_PYPDF:
        texto = limpar_texto(extrair_pypdf(pdf_path))
        score = avaliar_qualidade(texto)
        print(f"    Score de qualidade: {score:.2f}")
        if score >= LIMIAR:
            print(f"    ✔ pypdf suficiente.")
            return texto, tabelas, meta, "pypdf"
    else:
        print("    pypdf não disponível.")

    print(f"  Tentando motor 4: OCR (pytesseract)...")
    if HAS_OCR:
        texto = limpar_texto(extrair_ocr(pdf_path))
        score = avaliar_qualidade(texto)
        print(f"    Score de qualidade: {score:.2f}")
        if score > 0.1:
            print(f"    ✔ OCR concluído (score {score:.2f})")
            return texto, tabelas, meta, "OCR"
        else:
            print(f"    OCR retornou resultado insatisfatório.")
    else:
        print("    OCR não disponível (instale pytesseract e pdf2image).")

    # Se chegou aqui, pega o melhor que tiver
    if HAS_PYPDF:
        texto = limpar_texto(extrair_pypdf(pdf_path))
        return texto, tabelas, meta, "pypdf (baixa qualidade)"

    return "", [], meta, "nenhum"


# ──────────────────────────────────────────────────────
# Extrator inteligente de campos (heurísticas)
# ──────────────────────────────────────────────────────

# Padrões para documentos médicos, administrativos e jurídicos
PADROES = {
    "Paciente / Nome": [
        r'(?:paciente|nome do paciente|nome)\s*[:\-]?\s*([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ][^\n]{3,60})',
        r'(?:Sr\.|Sra\.|Dr\.|Dra\.)\s+([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ][a-záéíóúâêôãõüç ]{3,60})',
    ],
    "Data do Exame / Procedimento": [
        r'(?:data do exame|data de realiza[çc][aã]o|data|dt\.?)\s*[:\-]?\s*(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})',
        r'(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{4})',
    ],
    "Médico / Responsável": [
        r'(?:m[ée]dico responsável|m[ée]dico|dr\.|dra\.|crm\s*[:nº]?\s*\d+)\s*[:\-]?\s*([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ][^\n]{3,60})',
        r'(?:realizado por|executado por)\s*[:\-]?\s*([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ][^\n]{3,60})',
    ],
    "CRM": [
        r'CRM\s*[:\-nº°]?\s*(\d{3,8}(?:[/\-][A-Z]{2})?)',
    ],
    "Convênio / Operadora": [
        r'(?:conv[eê]nio|operadora|plano)\s*[:\-]?\s*([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ][^\n]{2,50})',
    ],
    "Número do Processo / Prontuário": [
        r'(?:prontu[aá]rio|n[uú]mero do processo|n[oº]\s*processo|processo)\s*[:\-nº°]?\s*([\d\.\-/]{3,30})',
        r'(?:n[uú]m(?:ero)?\.?|n[oº])\s*[:\-]?\s*([\d\.\-/]{5,20})',
    ],
    "CPF": [
        r'CPF\s*[:\-nº°]?\s*(\d{3}[\.\s]?\d{3}[\.\s]?\d{3}[\-\s]?\d{2})',
    ],
    "Data de Nascimento": [
        r'(?:nascimento|data de nasc(?:imento)?|dn\.?|nasc\.?)\s*[:\-]?\s*(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})',
    ],
    "Procedimento / Exame": [
        r'(?:procedimento|exame|servi[çc]o|t[íi]tulo|descri[çc][aã]o)\s*[:\-]?\s*([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ][^\n]{3,80})',
    ],
    "CID": [
        r'CID(?:\s*10)?\s*[:\-]?\s*([A-Z]\d{2}\.?\d*)',
    ],
    "Valor / Custo": [
        r'(?:valor|total|custo|pre[çc]o)\s*[:\-]?\s*R?\$?\s*([\d]{1,3}(?:[.,]\d{3})*(?:[.,]\d{2})?)',
    ],
    "Hospital / Clínica / Instituição": [
        r'(?:hospital|cl[íi]nica|instituto|centro m[ée]dico|unidade)\s+([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ][^\n]{2,60})',
    ],
    "Diagnóstico / Conclusão": [
        r'(?:diagn[óo]stico|conclus[aã]o|laudo|parecer|resultado)\s*[:\-]?\s*([^\n]{10,200})',
    ],
}


def extrair_campos_inteligentes(texto: str) -> dict:
    """
    Aplica padrões heurísticos para identificar campos relevantes no texto.
    Retorna dicionário campo -> lista de valores encontrados (sem duplicatas).
    """
    resultados = {}
    texto_busca = texto  # mantém case para nomes próprios
    texto_lower = texto.lower()

    for campo, padroes in PADROES.items():
        encontrados = set()
        for padrao in padroes:
            try:
                matches = re.findall(padrao, texto_busca, re.IGNORECASE | re.MULTILINE)
                for m in matches:
                    valor = m.strip() if isinstance(m, str) else m[0].strip()
                    if len(valor) > 1 and valor not in encontrados:
                        encontrados.add(valor)
            except Exception:
                continue
        if encontrados:
            # Limita a 3 ocorrências por campo para não poluir
            resultados[campo] = sorted(list(encontrados))[:3]

    return resultados


def formatar_tabela_texto(tabela: list) -> str:
    """Converte tabela (lista de listas) em texto formatado."""
    if not tabela:
        return ""
    linhas = []
    for linha in tabela:
        células = [str(c or '').strip() for c in linha]
        linhas.append(" | ".join(células))
    return '\n'.join(linhas)


# ──────────────────────────────────────────────────────
# Geração de TXT
# ──────────────────────────────────────────────────────

def gerar_txt(pdf_path: Path, texto: str, tabelas: list, meta: dict,
              campos: dict, motor: str, pasta_saida: Path) -> Path:
    """Gera arquivo TXT estruturado."""
    nome_base = pdf_path.stem
    arquivo_saida = pasta_saida / f"{nome_base}_CONVERTIDO.txt"

    agora = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    linhas = []
    linhas.append("=" * 70)
    linhas.append(f"  PDF CONVERTER - RELATÓRIO DE EXTRAÇÃO")
    linhas.append(f"  Arquivo: {pdf_path.name}")
    linhas.append(f"  Data/hora: {agora}")
    linhas.append(f"  Motor utilizado: {motor}")
    linhas.append("=" * 70)

    # Metadados
    if meta:
        linhas.append("\n[METADADOS DO ARQUIVO]")
        linhas.append(SEPARADOR)
        for k, v in meta.items():
            linhas.append(f"  {k}: {v}")

    # Campos extraídos inteligentemente
    if campos:
        linhas.append("\n[CAMPOS IDENTIFICADOS - EXTRAÇÃO INTELIGENTE]")
        linhas.append(SEPARADOR)
        for campo, valores in campos.items():
            for v in valores:
                linhas.append(f"  {campo}: {v}")

    # Tabelas
    if tabelas:
        linhas.append(f"\n[TABELAS ENCONTRADAS: {len(tabelas)}]")
        linhas.append(SEPARADOR)
        for (pag, idx, tabela) in tabelas:
            linhas.append(f"\n  Tabela {idx} - Página {pag}:")
            linhas.append(formatar_tabela_texto(tabela))

    # Texto completo
    linhas.append("\n[TEXTO EXTRAÍDO COMPLETO]")
    linhas.append(SEPARADOR)
    linhas.append(texto)

    conteudo = '\n'.join(linhas)

    with open(arquivo_saida, 'w', encoding='utf-8') as f:
        f.write(conteudo)

    return arquivo_saida


# ──────────────────────────────────────────────────────
# Geração de DOCX
# ──────────────────────────────────────────────────────

def gerar_docx(pdf_path: Path, texto: str, tabelas: list, meta: dict,
               campos: dict, motor: str, pasta_saida: Path) -> Path:
    """Gera arquivo DOCX formatado."""
    if not HAS_DOCX:
        print("  [DOCX] python-docx não disponível. Pulando geração DOCX.")
        return None

    nome_base = pdf_path.stem
    arquivo_saida = pasta_saida / f"{nome_base}_CONVERTIDO.docx"

    doc = Document()

    # Configurações de página
    section = doc.sections[0]
    section.page_width = int(21 * 914400 / 100 * 10)   # A4 aprox
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    agora = datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    # Título
    titulo = doc.add_heading(f'Conversão: {pdf_path.name}', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info da conversão
    p = doc.add_paragraph()
    p.add_run(f'Motor de extração: ').bold = False
    r = p.add_run(motor)
    r.bold = True
    p.add_run(f'    |    Data: {agora}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Metadados
    if meta:
        doc.add_heading('Metadados do Arquivo', level=2)
        tabela_meta = doc.add_table(rows=1, cols=2)
        tabela_meta.style = 'Table Grid'
        hdr = tabela_meta.rows[0].cells
        hdr[0].text = 'Campo'
        hdr[1].text = 'Valor'
        for k, v in meta.items():
            row = tabela_meta.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
        doc.add_paragraph()

    # Campos identificados
    if campos:
        doc.add_heading('Campos Identificados (Extração Inteligente)', level=2)
        tabela_campos = doc.add_table(rows=1, cols=2)
        tabela_campos.style = 'Table Grid'
        hdr2 = tabela_campos.rows[0].cells
        hdr2[0].text = 'Campo'
        hdr2[1].text = 'Valor(es) Encontrado(s)'
        for campo, valores in campos.items():
            row = tabela_campos.add_row().cells
            row[0].text = campo
            row[1].text = '\n'.join(valores)
        doc.add_paragraph()

    # Tabelas extraídas
    if tabelas:
        doc.add_heading(f'Tabelas Extraídas ({len(tabelas)})', level=2)
        for (pag, idx, tabela) in tabelas:
            doc.add_paragraph(f'Tabela {idx} - Página {pag}', style='Intense Quote')
            if tabela and len(tabela) > 0:
                try:
                    n_cols = max(len(r) for r in tabela if r)
                    if n_cols == 0:
                        continue
                    t = doc.add_table(rows=len(tabela), cols=n_cols)
                    t.style = 'Table Grid'
                    for i, linha in enumerate(tabela):
                        for j, célula in enumerate(linha):
                            if j < n_cols:
                                t.cell(i, j).text = str(célula or '')
                except Exception:
                    doc.add_paragraph(formatar_tabela_texto(tabela))
        doc.add_paragraph()

    # Texto completo
    doc.add_heading('Texto Extraído Completo', level=2)

    # Adiciona texto preservando estrutura de páginas
    for bloco in texto.split('\n\n'):
        bloco = bloco.strip()
        if not bloco:
            continue
        # Detecta marcador de página
        if re.match(r'^\[Página \d+\]$', bloco):
            p = doc.add_paragraph()
            r = p.add_run(bloco)
            r.bold = True
            r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        else:
            doc.add_paragraph(bloco)

    doc.save(str(arquivo_saida))
    return arquivo_saida


# ──────────────────────────────────────────────────────
# Leitura de caminho interativa
# ──────────────────────────────────────────────────────

def ler_caminho_usuario() -> Path:
    """
    Pergunta ao usuário o caminho do PDF ou pasta de PDFs.
    Aceita: arquivo único, pasta, caminhos Windows/WSL/Unix, com aspas.
    """
    print("\n" + "=" * 70)
    print("  PDF CONVERTER - Conversor e Extrator Inteligente")
    print("=" * 70)
    print("\nMotores disponíveis:")
    print(f"  {'✔' if HAS_PDFPLUMBER else '✗'} pdfplumber")
    print(f"  {'✔' if HAS_PYMUPDF else '✗'}  pymupdf")
    print(f"  {'✔' if HAS_PYPDF else '✗'}  pypdf")
    print(f"  {'✔' if HAS_OCR else '✗'}  OCR (pytesseract)")
    print(f"  {'✔' if HAS_DOCX else '✗'}  Geração DOCX")

    print("\nDigite o caminho do arquivo PDF ou da pasta com PDFs.")
    print("Exemplos:")
    print("  C:\\Users\\joao\\Documentos\\laudo.pdf")
    print("  C:\\Users\\joao\\OneDrive\\Área de Trabalho\\PDF_AMANDA")
    print("  /home/joao/documentos/laudo.pdf")
    print("  ./meus_pdfs/")
    print()

    while True:
        entrada = input("Caminho: ").strip()
        if not entrada:
            print("  Caminho vazio. Tente novamente.")
            continue

        caminho = normalizar_caminho(entrada)

        if caminho.is_file() and caminho.suffix.lower() == '.pdf':
            return caminho
        elif caminho.is_dir():
            pdfs = list(caminho.glob("*.pdf")) + list(caminho.glob("*.PDF"))
            if not pdfs:
                print(f"  Nenhum PDF encontrado em: {caminho}")
                print("  Tente outro caminho.")
                continue
            print(f"\n  {len(pdfs)} PDF(s) encontrado(s):")
            for p in pdfs:
                print(f"    - {p.name}")
            return caminho
        elif caminho.is_file() and caminho.suffix.lower() != '.pdf':
            print(f"  Arquivo '{caminho.name}' não é um PDF.")
            continue
        else:
            print(f"  Caminho não encontrado: {caminho}")
            print("  Verifique e tente novamente.")


# ──────────────────────────────────────────────────────
# Processamento principal
# ──────────────────────────────────────────────────────

def processar_pdf(pdf_path: Path, pasta_saida: Path) -> dict:
    """Processa um PDF e retorna relatório do resultado."""
    print(f"\n{'─' * 70}")
    print(f"  Processando: {pdf_path.name}")
    print(f"  Tamanho: {pdf_path.stat().st_size / 1024:.1f} KB")
    print(f"{'─' * 70}")

    resultado = {
        "arquivo": pdf_path.name,
        "sucesso": False,
        "motor": "",
        "chars": 0,
        "tabelas": 0,
        "campos": 0,
        "arquivos_gerados": []
    }

    # Extração em cascata
    texto, tabelas, meta, motor = extrair_texto_cascata(pdf_path)

    if not texto:
        print(f"  ✗ Não foi possível extrair texto de {pdf_path.name}")
        resultado["motor"] = motor
        return resultado

    # Extração inteligente de campos
    print(f"\n  Executando extrator inteligente de campos...")
    campos = extrair_campos_inteligentes(texto)
    print(f"  Campos identificados: {len(campos)}")

    # Gera TXT
    print(f"\n  Gerando TXT...")
    arq_txt = gerar_txt(pdf_path, texto, tabelas, meta, campos, motor, pasta_saida)
    print(f"  ✔ TXT salvo: {arq_txt.name}")
    resultado["arquivos_gerados"].append(str(arq_txt))

    # Gera DOCX
    print(f"  Gerando DOCX...")
    arq_docx = gerar_docx(pdf_path, texto, tabelas, meta, campos, motor, pasta_saida)
    if arq_docx:
        print(f"  ✔ DOCX salvo: {arq_docx.name}")
        resultado["arquivos_gerados"].append(str(arq_docx))

    resultado.update({
        "sucesso": True,
        "motor": motor,
        "chars": len(texto),
        "tabelas": len(tabelas),
        "campos": len(campos),
    })

    # Exibe prévia dos campos
    if campos:
        print(f"\n  Campos identificados:")
        for campo, valores in campos.items():
            print(f"    {campo}: {valores[0][:80]}")

    return resultado


def main():
    caminho = ler_caminho_usuario()

    # Define lista de PDFs a processar
    if caminho.is_file():
        pdfs = [caminho]
        pasta_saida = caminho.parent / "CONVERTIDOS"
    else:
        pdfs = sorted(list(caminho.glob("*.pdf")) + list(caminho.glob("*.PDF")))
        pasta_saida = caminho / "CONVERTIDOS"

    pasta_saida.mkdir(exist_ok=True)
    print(f"\n  Pasta de saída: {pasta_saida}")

    # Processamento
    relatorio = []
    for pdf in pdfs:
        r = processar_pdf(pdf, pasta_saida)
        relatorio.append(r)

    # Resumo final
    print(f"\n{'=' * 70}")
    print(f"  RESUMO FINAL")
    print(f"{'=' * 70}")
    sucessos = sum(1 for r in relatorio if r["sucesso"])
    print(f"  Processados: {len(pdfs)} | Sucesso: {sucessos} | Falhas: {len(pdfs) - sucessos}")
    print()
    for r in relatorio:
        status = "✔" if r["sucesso"] else "✗"
        print(f"  {status} {r['arquivo']}")
        if r["sucesso"]:
            print(f"      Motor: {r['motor']} | {r['chars']:,} chars | "
                  f"{r['tabelas']} tabelas | {r['campos']} campos extraídos")
    print(f"\n  Arquivos salvos em: {pasta_saida}")
    print()


if __name__ == "__main__":
    main()
